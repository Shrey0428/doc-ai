import streamlit as st
import tempfile
import fitz  # PyMuPDF
from PIL import Image
from google.cloud import vision
from google.oauth2 import service_account
import openai
from docx import Document
import os

# === Streamlit Config ===
st.set_page_config(page_title="📄 Legal Document Decoder", layout="wide")
st.title("📤 Upload Legal PDF → 🔍 OCR → 🤖 GPT-4 → 📄 Word Export")

# === Load API Key from Streamlit Secrets ===
openai.api_key = st.secrets["OPENAI_API_KEY"]

# === Google Cloud Vision Auth ===
CREDENTIAL_PATH = "gcloud_key.json"
credentials = service_account.Credentials.from_service_account_file(CREDENTIAL_PATH)
vision_client = vision.ImageAnnotatorClient(credentials=credentials)

# === Upload PDF File ===
uploaded_file = st.file_uploader("Upload a scanned legal PDF", type=["pdf"])

if uploaded_file:
    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.read())
        pdf_path = tmp_file.name

    st.success("✅ PDF uploaded. Performing OCR...")

    # === Extract Text from PDF Pages ===
    doc = fitz.open(pdf_path)
    full_text = ""

    for page in doc:
        pix = page.get_pixmap(dpi=300)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
            img_path = tmp_img.name
            pix.save(img_path)

            with open(img_path, "rb") as img_file:
                content = img_file.read()
                image = vision.Image(content=content)
                response = vision_client.document_text_detection(image=image)
                full_text += response.full_text_annotation.text + "\n"

    # === Display OCR Result ===
    st.text_area("📃 Extracted OCR Text", full_text, height=300)

    # === Analyze with GPT-4 ===
    if st.button("🧠 Analyze with GPT-4"):
        st.info("Running GPT-4 legal analysis...")

        prompt = f"""
You are a legal assistant. Extract structured legal information from this scanned document.

Return a JSON object with:
- Agreement Type
- Parties Involved
- Effective Date
- Jurisdiction
- Monetary Amounts
- Duration
- Key Clauses

OCR Document Text:
{full_text}
"""

        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You extract structured legal data from OCR'd documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2
        )

        result = response.choices[0].message.content.strip()
        st.code(result, language="json")

        # === Export to Word Document ===
        docx = Document()
        docx.add_heading("Legal Document Summary", level=1)
        docx.add_paragraph(result)

        word_path = os.path.join(tempfile.gettempdir(), "Legal_Summary.docx")
        docx.save(word_path)

        with open(word_path, "rb") as f:
            st.download_button("⬇️ Download Word Document", f, file_name="Legal_Summary.docx")
